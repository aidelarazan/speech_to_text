#!/usr/bin/env python
# coding: utf-8

# # Speech to Text
# ## Angelique I. Delarazan

# ### Transcribes audio recordings to text and outputs it as a Word Doc file.
# 
# Original code by Anthony Zhang:
# Zhang, A. (2017). Speech Recognition (Version 3.8) [Software]. Available from https://github.com/Uberi/speech_recognition#readme.
# 
# Requirements:
# - Python 2.6, 2.7, or 3.3+ (required)
# - Installation of SpeechRecognition
# - Installation of Docx
# - Audio recording in .wav format (mono)

# ### Import Packages

# OS
import os
from os import path

# Speech Recognition
import speech_recognition as sr

# Docx
import docx
from docx import Document
from docx.shared import Inches

# Read/write wav files
import wave

# with-statement context
import contextlib


# ### Set Up

# Define subject or a list of subjects (or name of audio file(s))
subject = 315

# Set directory (change paths/file names accordingly)
audio_file = path.join(path.dirname('/Users/aidelarazan/Downloads/Speech_To_Text/'), ("%s_recall.wav" %(subject)))
file_name = os.path.basename(audio_file)
name = os.path.basename(audio_file)

# Use the audio file as the audio source
r = sr.Recognizer()

# Create a new word doc file
document = Document()

# Read in audio file
wav = ("%s_recall.wav" %(subject))


# ### Audio to Text Transcription

# Obtain duration of audio file
with contextlib.closing(wave.open(wav,'r')) as f:
    frames = f.getnframes()
    rate = f.getframerate()
    time = frames / float(rate)
    t = str(time)
    p = document.add_paragraph([t])
    document.save(("%s_recall.docx" %(subject))) # Add duration to word doc file; change output path accordingly

# Transcribing...    
count = -20
while count < time:
    count += 15  # Transcribes every 15 seconds; this is the same as count = count + 1
    r = sr.Recognizer()
    with sr.AudioFile(audio_file) as source:
        audio = r.record(source, offset=count, duration=15)  # read the entire audio file   
    try:
    # Recognize speech using Google Speech Recognition; current code uses the default Speech Recognition API key
        # to use another API key, use `r.recognize_google(audio, key="GOOGLE_SPEECH_RECOGNITION_API_KEY")`
        # instead of `r.recognize_google(audio)` 
        transcription =  r.recognize_google(audio)
        #print(transcription) # Comment line out if do not want to see transcription outputs
        p = document.add_paragraph([transcription])
        document.save(("%s_recall.docx" %(subject)))
    # Errors     
    except sr.UnknownValueError:
        transcription = "Google Speech Recognition could not understand audio"
    except sr.RequestError as e:
        transcription = "Could not request results from Google Speech Recognition service; {0}".format(e)

